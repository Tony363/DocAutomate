#!/usr/bin/env python3
"""
Document Ingestion Pipeline
Supports multiple document formats and extracts structured content
"""

import hashlib
import json
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import yaml
import argparse
import mimetypes

# Document parser imports (install via pip if needed)
try:
    import pypdf
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import markdown
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

try:
    from bs4 import BeautifulSoup
    HAS_HTML = True
except ImportError:
    HAS_HTML = False


class DocumentIngestion:
    """Main document ingestion pipeline"""
    
    def __init__(self, output_dir: str = "docs/source"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def ingest_document(self, file_path: str) -> Dict[str, Any]:
        """
        Ingest a document and extract structured content
        Returns document metadata and extracted content
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Generate document ID from content hash
        doc_id = self._generate_doc_id(file_path)
        
        # Detect document format
        doc_format = self._detect_format(file_path)
        
        # Extract content based on format
        content, structure, fidelity = self._extract_content(file_path, doc_format)
        
        # Create section index
        sections = self._create_section_index(structure)
        
        # Generate metadata
        metadata = {
            "doc_id": doc_id,
            "original_path": str(file_path),
            "format": doc_format,
            "ingestion_date": datetime.now().isoformat(),
            "fidelity_score": fidelity,
            "file_size": file_path.stat().st_size,
            "sections_count": len(sections)
        }
        
        # Save extracted content
        output_path = self._save_content(doc_id, content, structure, metadata)
        
        return {
            "doc_id": doc_id,
            "format": doc_format,
            "fidelity": fidelity,
            "sections": sections,
            "output_path": str(output_path),
            "metadata": metadata
        }
    
    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate content-addressed document ID"""
        hasher = hashlib.sha256()
        with open(file_path, 'rb') as f:
            while chunk := f.read(8192):
                hasher.update(chunk)
        
        # Include filename for readability
        name_part = file_path.stem[:20].replace(" ", "_").lower()
        hash_part = hasher.hexdigest()[:12]
        return f"{name_part}_{hash_part}"
    
    def _detect_format(self, file_path: Path) -> str:
        """Detect document format from extension and MIME type"""
        ext = file_path.suffix.lower()
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        format_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.html': 'html',
            '.htm': 'html',
            '.txt': 'plaintext'
        }
        
        return format_map.get(ext, 'plaintext')
    
    def _extract_content(self, file_path: Path, doc_format: str) -> Tuple[str, Dict, float]:
        """
        Extract content from document
        Returns: (plain_text, structure, fidelity_score)
        """
        if doc_format == 'pdf':
            return self._extract_pdf(file_path)
        elif doc_format == 'docx':
            return self._extract_docx(file_path)
        elif doc_format == 'markdown':
            return self._extract_markdown(file_path)
        elif doc_format == 'html':
            return self._extract_html(file_path)
        else:
            return self._extract_plaintext(file_path)
    
    def _extract_pdf(self, file_path: Path) -> Tuple[str, Dict, float]:
        """Extract content from PDF"""
        if not HAS_PDF:
            print("Warning: pypdf not installed, using fallback text extraction")
            return self._extract_plaintext(file_path)
        
        try:
            reader = pypdf.PdfReader(str(file_path))
            text_parts = []
            structure = {"pages": [], "headings": []}
            
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                text_parts.append(page_text)
                structure["pages"].append({
                    "page_num": i + 1,
                    "text_length": len(page_text)
                })
            
            content = "\n\n".join(text_parts)
            
            # Extract headings heuristically
            for line in content.split('\n'):
                if line.isupper() and len(line) > 5:
                    structure["headings"].append(line)
            
            fidelity = 0.8  # PDF extraction is generally good
            return content, structure, fidelity
            
        except Exception as e:
            print(f"PDF extraction error: {e}")
            return self._extract_plaintext(file_path)
    
    def _extract_docx(self, file_path: Path) -> Tuple[str, Dict, float]:
        """Extract content from DOCX"""
        if not HAS_DOCX:
            print("Warning: python-docx not installed, using fallback")
            return self._extract_plaintext(file_path)
        
        try:
            doc = DocxDocument(str(file_path))
            text_parts = []
            structure = {"paragraphs": [], "headings": [], "tables": []}
            
            for para in doc.paragraphs:
                text_parts.append(para.text)
                if para.style.name.startswith('Heading'):
                    structure["headings"].append({
                        "level": para.style.name,
                        "text": para.text
                    })
                structure["paragraphs"].append({
                    "style": para.style.name,
                    "length": len(para.text)
                })
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    rows.append(cells)
                structure["tables"].append({
                    "index": i,
                    "rows": len(table.rows),
                    "cols": len(table.columns)
                })
            
            content = "\n".join(text_parts)
            fidelity = 0.9  # DOCX extraction is very good
            return content, structure, fidelity
            
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return self._extract_plaintext(file_path)
    
    def _extract_markdown(self, file_path: Path) -> Tuple[str, Dict, float]:
        """Extract content from Markdown"""
        content = file_path.read_text(encoding='utf-8')
        structure = {"headings": [], "code_blocks": [], "links": []}
        
        # Extract headings
        for line in content.split('\n'):
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                structure["headings"].append({
                    "level": level,
                    "text": text
                })
        
        # Extract code blocks
        import re
        code_blocks = re.findall(r'```[\s\S]*?```', content)
        structure["code_blocks"] = [{"index": i, "length": len(block)} 
                                   for i, block in enumerate(code_blocks)]
        
        # Extract links
        links = re.findall(r'\[([^\]]+)\]\(([^\)]+)\)', content)
        structure["links"] = [{"text": text, "url": url} for text, url in links]
        
        fidelity = 1.0  # Markdown is perfectly preserved
        return content, structure, fidelity
    
    def _extract_html(self, file_path: Path) -> Tuple[str, Dict, float]:
        """Extract content from HTML"""
        if not HAS_HTML:
            print("Warning: beautifulsoup4 not installed, using fallback")
            return self._extract_plaintext(file_path)
        
        try:
            html_content = file_path.read_text(encoding='utf-8')
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for element in soup(['script', 'style']):
                element.decompose()
            
            # Extract text
            content = soup.get_text(separator='\n', strip=True)
            
            # Extract structure
            structure = {
                "headings": [],
                "links": [],
                "images": []
            }
            
            for tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                for heading in soup.find_all(tag):
                    structure["headings"].append({
                        "level": tag,
                        "text": heading.get_text(strip=True),
                        "id": heading.get('id', '')
                    })
            
            for link in soup.find_all('a'):
                structure["links"].append({
                    "text": link.get_text(strip=True),
                    "href": link.get('href', '')
                })
            
            for img in soup.find_all('img'):
                structure["images"].append({
                    "alt": img.get('alt', ''),
                    "src": img.get('src', '')
                })
            
            fidelity = 0.85  # HTML extraction is good
            return content, structure, fidelity
            
        except Exception as e:
            print(f"HTML extraction error: {e}")
            return self._extract_plaintext(file_path)
    
    def _extract_plaintext(self, file_path: Path) -> Tuple[str, Dict, float]:
        """Fallback plain text extraction"""
        content = file_path.read_text(encoding='utf-8', errors='ignore')
        structure = {
            "lines": len(content.split('\n')),
            "words": len(content.split()),
            "chars": len(content)
        }
        fidelity = 0.5  # Plain text loses structure
        return content, structure, fidelity
    
    def _create_section_index(self, structure: Dict) -> List[Dict]:
        """Create searchable section index from structure"""
        sections = []
        
        # Process headings if available
        if 'headings' in structure:
            for i, heading in enumerate(structure['headings']):
                section = {
                    "index": i,
                    "type": "heading",
                    "level": heading.get('level', 1),
                    "title": heading.get('text', ''),
                    "path": f"/section_{i}"
                }
                sections.append(section)
        
        # Process tables if available
        if 'tables' in structure:
            for i, table in enumerate(structure['tables']):
                section = {
                    "index": len(sections) + i,
                    "type": "table",
                    "title": f"Table {i+1}",
                    "path": f"/table_{i}",
                    "dimensions": f"{table.get('rows')}x{table.get('cols')}"
                }
                sections.append(section)
        
        # Process code blocks if available
        if 'code_blocks' in structure:
            for i, block in enumerate(structure['code_blocks']):
                section = {
                    "index": len(sections) + i,
                    "type": "code",
                    "title": f"Code Block {i+1}",
                    "path": f"/code_{i}"
                }
                sections.append(section)
        
        return sections
    
    def _save_content(self, doc_id: str, content: str, structure: Dict, 
                     metadata: Dict) -> Path:
        """Save extracted content and metadata"""
        doc_dir = self.output_dir / doc_id
        doc_dir.mkdir(parents=True, exist_ok=True)
        
        # Save plain text content
        content_file = doc_dir / "content.txt"
        content_file.write_text(content, encoding='utf-8')
        
        # Save structure as JSON
        structure_file = doc_dir / "structure.json"
        with open(structure_file, 'w') as f:
            json.dump(structure, f, indent=2)
        
        # Save metadata as YAML
        metadata_file = doc_dir / "metadata.yaml"
        with open(metadata_file, 'w') as f:
            yaml.dump(metadata, f, default_flow_style=False)
        
        # Create ingestion report
        report = {
            "doc_id": doc_id,
            "status": "success",
            "files_created": [
                str(content_file),
                str(structure_file),
                str(metadata_file)
            ],
            "timestamp": datetime.now().isoformat()
        }
        
        report_file = doc_dir / "ingestion_report.json"
        with open(report_file, 'w') as f:
            json.dump(report, f, indent=2)
        
        return doc_dir


def main():
    """CLI interface for document ingestion"""
    parser = argparse.ArgumentParser(description="Document Ingestion Pipeline")
    parser.add_argument("document", help="Path to document to ingest")
    parser.add_argument("--output-dir", default="docs/source",
                       help="Output directory for extracted content")
    parser.add_argument("--verbose", action="store_true",
                       help="Verbose output")
    
    args = parser.parse_args()
    
    # Check dependencies
    if not any([HAS_PDF, HAS_DOCX, HAS_MARKDOWN, HAS_HTML]):
        print("Warning: No specialized parsers installed.")
        print("Install with: pip install pypdf python-docx markdown beautifulsoup4")
    
    # Run ingestion
    pipeline = DocumentIngestion(output_dir=args.output_dir)
    
    try:
        result = pipeline.ingest_document(args.document)
        
        if args.verbose:
            print(f"Document ingested successfully!")
            print(f"Document ID: {result['doc_id']}")
            print(f"Format: {result['format']}")
            print(f"Fidelity: {result['fidelity']:.2f}")
            print(f"Sections: {len(result['sections'])}")
            print(f"Output: {result['output_path']}")
        else:
            print(json.dumps(result, indent=2))
            
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()